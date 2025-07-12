import pandas as pd  # type: ignore
import tkinter as tk
from tkinter import filedialog, messagebox
import datetime
import locale
from docx import Document
from docx.shared import Cm, Pt, Inches
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define o locale em português para exibir meses corretamente
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error:
    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')

def set_cell_width(cell, width_cm):
    width_twips = int(width_cm * 567)
    tcPr = cell._tc.get_or_add_tcPr()
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    cell.width = Cm(width_cm)

def set_table_width(table, width_cm):
    width_twips = int(width_cm * 567)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def aplicar_fonte_tamanho_10(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)

class BorderoApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Borderô de Direitos Autorais")
        self.root.geometry("600x400")
        self.arquivo_excel = None
        self.df = None
        self.data_inicio_filtro = "01/01/2023"
        self.data_fim_filtro = "31/12/2023"
        self.codigo_produto = tk.StringVar()
        self.percentual_direitos = tk.DoubleVar(value=10.0)
        self.criar_interface()

    def criar_interface(self):
        frame = tk.Frame(self.root, padx=20, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)

        tk.Button(frame, text="Selecionar Planilha Excel", command=self.selecionar_arquivo).grid(row=0, column=0, columnspan=2, pady=10, sticky="ew")

        tk.Label(frame, text="Código do Produto:").grid(row=1, column=0, pady=10, sticky="w")
        tk.Entry(frame, textvariable=self.codigo_produto).grid(row=1, column=1, pady=10, sticky="ew")

        tk.Label(frame, text="Percentual de Direitos (%):").grid(row=2, column=0, pady=10, sticky="w")
        tk.Entry(frame, textvariable=self.percentual_direitos).grid(row=2, column=1, pady=10, sticky="ew")

        tk.Button(frame, text="Gerar Recibo", command=self.gerar_recibo).grid(row=3, column=0, pady=20, padx=5, sticky="ew")
        tk.Button(frame, text="Gerar Lista Detalhada", command=self.gerar_lista_detalhada).grid(row=3, column=1, pady=20, padx=5, sticky="ew")

        self.status_label = tk.Label(frame, text="Aguardando arquivo...", fg="blue")
        self.status_label.grid(row=4, column=0, columnspan=2, pady=10)

        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=2)

    def selecionar_arquivo(self):
        arquivo = filedialog.askopenfilename(title="Selecione o arquivo Excel", filetypes=[("Excel files", "*.xlsx;*.xls")])
        if arquivo:
            self.arquivo_excel = arquivo
            try:
                # Lê A1 para extrair as datas do filtro
                a1_texto = pd.read_excel(arquivo, header=None).iloc[0, 0]
                padrao = r'Data de Cadastro de (\d{2}/\d{2}/\d{4}) até (\d{2}/\d{2}/\d{4})'
                match = re.search(padrao, str(a1_texto))
                if match:
                    self.data_inicio_filtro = match.group(1)
                    self.data_fim_filtro = match.group(2)

                self.df = pd.read_excel(arquivo, header=3)
                self.df = self.df.iloc[:-1]
                self.df.columns = self.df.columns.str.strip()
                self.df["Preço Venda Total (R$)"] = pd.to_numeric(self.df["Preço Venda Total (R$)"], errors='coerce')
                self.status_label.config(text=f"Arquivo carregado: {arquivo.split('/')[-1]}", fg="green")
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao ler o arquivo: {str(e)}")
                self.status_label.config(text="Erro ao carregar arquivo", fg="red")

    def filtrar_dados(self):
        if self.df is None:
            messagebox.showwarning("Aviso", "Nenhum arquivo Excel carregado!")
            return None

        codigo = self.codigo_produto.get()
        if not codigo:
            messagebox.showwarning("Aviso", "Digite o código do produto!")
            return None

        try:
            colunas_disponiveis = self.df.columns.str.strip().tolist()
            if "Código Prod." not in colunas_disponiveis:
                messagebox.showerror("Erro", f"Coluna 'Código Prod.' não encontrada!\nColunas disponíveis: {colunas_disponiveis}")
                return None

            codigos_df = self.df["Código Prod."].astype(float).astype(int).astype(str)
            codigo_str = str(int(float(codigo)))

            df_filtrado = self.df[codigos_df == codigo_str]

            if df_filtrado.empty:
                codigos_unicos = codigos_df.unique()
                messagebox.showinfo("Informação", f"Nenhum registro encontrado para o código '{codigo_str}'.\n\nCódigos disponíveis:\n{', '.join(codigos_unicos)}")
                return None

            return df_filtrado

        except KeyError as e:
            messagebox.showerror("Erro", f"Erro ao acessar os dados: {e}")
            return None

    def gerar_recibo(self):
        df_filtrado = self.filtrar_dados()
        if df_filtrado is None:
            return

        try:
            titulo_original = df_filtrado.iloc[0]["Descrição"] if "Descrição" in df_filtrado.columns else "TÍTULO NÃO ENCONTRADO"
            titulo_produto = titulo_original.split(" - ", 1)[1] if " - " in titulo_original else titulo_original
            total_vendas = df_filtrado["Quantidade"].sum()
            valor_total = df_filtrado["Preço Venda Total (R$)"].sum() if "Preço Venda Total (R$)" in df_filtrado.columns else 0
            percentual = self.percentual_direitos.get() / 100
            direitos_autorais = round(valor_total * percentual, 2)
            percentual_str = f"{self.percentual_direitos.get():.2f}"
            autor = df_filtrado.iloc[0]["Fornecedor"] if "Fornecedor" in df_filtrado.columns else "AUTOR NÃO ENCONTRADO"

            substituicoes = {
                '[produto]': titulo_produto,
                '[direitos]': f"{direitos_autorais:.2f}",
                '[direitosExtenso]': self.numero_por_extenso(direitos_autorais),
                '[dataInicio]': self.data_inicio_filtro,
                '[dataFim]': self.data_fim_filtro,
                '[dataRecibo]': datetime.datetime.now().strftime("%d de %B de %Y").capitalize(),
                '[autor]': autor,
                '[vendas]': str(total_vendas),
                '[valorTotal]': f'{valor_total:.2f}',
                '[percentual]': percentual_str
            }

            try:
                document = Document('modelo-recibo.docx')
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao carregar modelo: {str(e)}")
                return

            for p in document.paragraphs:
                for key, val in substituicoes.items():
                    if key in p.text:
                        inline = p.runs
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                inline[i].text = inline[i].text.replace(key, val)

            nome_produto_limpo = "".join(c for c in titulo_produto if c.isalnum() or c in (' ', '-', '_')).strip()
            nome_arquivo_final = f"Recibo - {nome_produto_limpo}.docx"
            arquivo_saida = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=nome_arquivo_final,
                filetypes=[("Documentos Word", "*.docx")]
            )
            if arquivo_saida:
                document.save(arquivo_saida)
                messagebox.showinfo("Sucesso", f"Recibo salvo com sucesso em {arquivo_saida}")

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar recibo: {str(e)}")

    def gerar_lista_detalhada(self):
        df_filtrado = self.filtrar_dados()
        if df_filtrado is None:
            return
        try:
            document = Document('modelo-lista.docx')
            titulo_original = df_filtrado.iloc[0].get("Descrição", "TÍTULO NÃO ENCONTRADO")
            titulo_produto = titulo_original.split(" - ", 1)[1] if " - " in titulo_original else titulo_original
            autor = df_filtrado.iloc[0].get("Fornecedor", "AUTOR NÃO ENCONTRADO")
            total_valor = df_filtrado["Preço Venda Total (R$)"].sum()
            total_quantidade = df_filtrado["Quantidade"].sum()
            percentual = self.percentual_direitos.get()

            substituicoes = {
                '[autor]': autor,
                '[dataInicio]': self.data_inicio_filtro,
                '[dataFim]': self.data_fim_filtro,
                '[codProd]': self.codigo_produto.get(),
                '[produto]': titulo_produto,
                '[qtdTotal]': str(total_quantidade),
                '[valorTotal]': f"{total_valor:.2f}",
                '[percentual]': f"{percentual:.2f}",
                '[valor_direitos]': f"{total_valor * percentual / 100:.2f}",
            }

            self.replace_all_paragraphs(document.paragraphs, substituicoes)
            for marcador in ['[dataVendas]', '[codigoVendas]', '[codNfe]', '[cliente]', '[precoVenda]', '[qtd]', '[precoVendaTotal]']:
                self.remover_linha_com_marcador(document, marcador)

            idx = None
            for i, p in enumerate(document.paragraphs):
                if "[tabela]" in p.text:
                    p.text = ""
                    idx = i
                    break
            if idx is None:
                messagebox.showerror("Erro", "Marcador [tabela] não encontrado no modelo.")
                return

            table = document.add_table(rows=1, cols=8)
            table.autofit = False
            set_table_width(table, 17.9)
            col_widths = [2.4, 1.7, 1.3, 5.0, 2.0, 1.2, 2.2, 2.1]
            headers = ['Data', 'Pedido', 'NF', 'Cliente', 'Pr. bruto', 'Qtd.', 'Pr. líquido', 'Total item']
            hdr_cells = table.rows[0].cells
            for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
                cell.text = header
                aplicar_fonte_tamanho_10(cell)
                set_cell_width(cell, col_widths[i])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                tcPr = cell._tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), 'center')
                tcPr.append(tcVAlign)

            linha_preta = table.add_row().cells
            for i in range(len(linha_preta)):
                linha_preta[i].text = ""
                aplicar_fonte_tamanho_10(linha_preta[i])
                set_cell_width(linha_preta[i], col_widths[i])
                tcPr = linha_preta[i]._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)

            for _, row in df_filtrado.iterrows():
                data = pd.to_datetime(row.get("Data", "")).strftime("%d/%m/%Y") if pd.notna(row.get("Data", "")) else ""
                pedido_val = row.get("Código Venda", "")
                pedido = str(int(pedido_val)) if pd.notna(pedido_val) and pedido_val != "" else ""
                nf_val = row.get("Cód. NFe/NFSe", "")
                nf = str(int(nf_val)) if pd.notna(nf_val) and nf_val != "" else ""
                cliente = str(row.get("Cliente", ""))
                preco = float(row.get("Preço Venda (R$)", 0.0))
                qtd = int(row.get("Quantidade", 0))
                total = float(row.get("Preço Venda Total (R$)", 0.0))
                preco_liquido = total / qtd if qtd else 0.0
                dados = [data, pedido, nf, cliente, f"{preco:.2f}", str(qtd), f"{preco_liquido:.2f}", f"{total:.2f}"]

                row_cells = table.add_row().cells
                for i, valor in enumerate(dados):
                    row_cells[i].text = valor
                    aplicar_fonte_tamanho_10(row_cells[i])
                    set_cell_width(row_cells[i], col_widths[i])
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tcPr = row_cells[i]._tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), 'center')
                    tcPr.append(tcVAlign)

                blank_row = table.add_row().cells
                for i in range(len(blank_row)):
                    blank_row[i].text = ""
                    aplicar_fonte_tamanho_10(blank_row[i])
                    set_cell_width(blank_row[i], col_widths[i])

            linha_preta_final = table.add_row().cells
            for i in range(len(linha_preta_final)):
                linha_preta_final[i].text = ""
                aplicar_fonte_tamanho_10(linha_preta_final[i])
                set_cell_width(linha_preta_final[i], col_widths[i])
                tcPr = linha_preta_final[i]._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)

            document._element.body.insert(idx + 1, table._element)

            nome_produto_limpo = "".join(c for c in titulo_produto if c.isalnum() or c in (' ', '-', '_')).strip()
            nome_arquivo_final = f"Lista Detalhada - {nome_produto_limpo}.docx"
            arquivo_saida = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=nome_arquivo_final, filetypes=[("Documentos Word", "*.docx")])
            if arquivo_saida:
                document.save(arquivo_saida)
                messagebox.showinfo("Sucesso", f"Lista detalhada salva com sucesso em {arquivo_saida}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar lista detalhada: {str(e)}")

    def replace_all_paragraphs(self, paragraphs, substitutions):
        for p in paragraphs:
            for key, val in substitutions.items():
                if key in p.text:
                    inline = p.runs
                    new_text = p.text.replace(key, val)
                    for i in range(len(inline)):
                        inline[i].text = ""
                    if inline:
                        inline[0].text = new_text

    def remover_linha_com_marcador(self, document, marcador):
        for p in document.paragraphs:
            if marcador in p.text:
                p.clear()

    def numero_por_extenso(self, valor):
        try:
            from num2words import num2words
            parte_inteira = int(valor)
            parte_decimal = int(round((valor - parte_inteira) * 100))
            extenso = num2words(parte_inteira, lang='pt_BR') + ' reais'
            if parte_decimal > 0:
                extenso += ' e ' + num2words(parte_decimal, lang='pt_BR') + ' centavos'
            return extenso.capitalize()
        except ImportError:
            return f"{valor:.2f} Reais"
        except Exception as e:
            return f"{valor:.2f} Reais (Erro ao converter para extenso: {e})"
        
if __name__ == "__main__":
    root = tk.Tk()
    app = BorderoApp(root)
    root.mainloop()